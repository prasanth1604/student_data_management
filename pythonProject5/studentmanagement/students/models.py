from django.db import models
import pandas as pd
from io import BytesIO
from docx import Document

class Student(models.Model):
    roll_no = models.CharField(max_length=8, unique=True)
    name = models.CharField(max_length=30)
    gender = models.CharField(max_length=10, default='NA')
    parent = models.CharField(max_length=30, default='NA')
    address = models.CharField(max_length=50, default='NA')


    def __str__(self):
        return self.roll_no

    @staticmethod
    def export_to_excel():
        # Gather all student data along with their assignments
        data = []
        for student in Student.objects.all():
            try:
                subject = Subject.objects.get(rollNo=student.roll_no)
                data.append({
                    "Roll No": student.roll_no,
                    "Name": student.name,
                    "Gender": student.gender,
                    "Parent": student.parent,
                    "Address": student.address,
                    "Subject Code": subject.subjectCode,
                    "Assignment 1": subject.assignment1,
                    "Assignment 2": subject.assignment2,
                    "Total": subject.total,
                })
            except Subject.DoesNotExist:
                data.append({
                    "Roll No": student.roll_no,
                    "Name": student.name,
                    "Gender": student.gender,
                    "Parent": student.parent,
                    "Address": student.address,
                    "Assignment 1": "N/A",
                    "Assignment 2": "N/A",
                    "Total": "N/A",
                })

        # Create a DataFrame using pandas
        df = pd.DataFrame(data)

        # Create an Excel writer using BytesIO as a buffer
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Students')

        # Seek to the beginning of the stream
        buffer.seek(0)
        return buffer

#Similarly to export to word document
    @staticmethod
    def export_to_word():
        # Gather all student data along with their assignments
        data = []
        for student in Student.objects.all():
            try:
                subject = Subject.objects.get(rollNo=student.roll_no)
                data.append({
                    "Roll No": student.roll_no,
                    "Name": student.name,
                    "Gender": student.gender,
                    "Parent": student.parent,
                    "Address": student.address,
                    "Subject Code": subject.subjectCode,
                    "Assignment 1": subject.assignment1,
                    "Assignment 2": subject.assignment2,
                    "Total": subject.total,
                })
            except Subject.DoesNotExist:
                data.append({
                    "Roll No": student.roll_no,
                    "Name": student.name,
                    "Gender": student.gender,
                    "Parent": student.parent,
                    "Address": student.address,
                    "Assignment 1": "N/A",
                    "Assignment 2": "N/A",
                    "Total": "N/A",
                })

        # Create a Word document
        doc = Document()
        doc.add_heading('Student Data', level=1)

        # Add table to the document
        table = doc.add_table(rows=1, cols=len(data[0]))
        hdr_cells = table.rows[0].cells
        headers = ["Roll No", "Name", "Gender", "Parent", "Address", "Subject Code", "Assignment 1", "Assignment 2",
                   "Total"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        for student_data in data:
            row_cells = table.add_row().cells
            for i, (key, value) in enumerate(student_data.items()):
                row_cells[i].text = str(value)

        # Save the document to a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


class Faculty(models.Model):
    fId = models.CharField(max_length=20, unique=True)
    name = models.CharField(max_length=30)
    subjectName = models.CharField(max_length=50)
    subjectCode = models.CharField(max_length=20)
    password = models.CharField(max_length=100)

    def __str__(self):
        return self.name


class Subject(models.Model):
    rollNo = models.CharField(max_length=8, unique=True)
    name = models.CharField(max_length=30, unique=False)
    subjectCode = models.CharField(max_length=20, unique=False)
    assignment1 = models.FloatField(default=0.00)
    assignment2 = models.FloatField(default=0.00)
    total = models.FloatField(default=0.00)

    def __str__(self):
        return self.subjectCode + " " + self.rollNo




