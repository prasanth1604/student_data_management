from django.shortcuts import render, redirect
from .models import Student, Faculty
from django.http import HttpResponse, HttpResponseForbidden
import pandas as pd
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom


def login(request):
    error = ""
    if request.method == 'POST':
        roll = request.POST['Roll_No']
        global st_id  # st_id is the id of the logged in student used in other templates (until logged out)
        st_id = roll
        error = ""
        pwd = request.POST['password']
        if Student.objects.filter(roll_no=roll, password=pwd):
            error += "success"

            return dashboard(request)
        else:
            error += "Try again"

    return render(request, 'login.html', {'error': error})


def loginAsFaculty(request):
    error = ""
    if request.method == 'POST':
        Faculty_Id = request.POST['Faculty_Id']
        global f_id  # f_id is the id of the logged in Faculty used in other templates (until logged out)
        f_id = Faculty_Id
        pwd = request.POST['password']
        if Faculty.objects.filter(fId=Faculty_Id, password=pwd):
            error = "success"
            s = "abc"
            return facultydashboard(request)
        else:
            error += "Try again"
    return render(request, 'loginAsfaculty.html', {'error': error})


def facultydashboard(request):
    if request.method == 'POST':
        roll = request.POST.get('rol')
        sCode = request.POST.get('subCode')
        A1 = request.POST.get('a1')
        A2 = request.POST.get('a2')
        facultyitem = Faculty.objects.get(fId=f_id)

        if roll != None and Student.objects.get(roll_no=roll, subjectCode=sCode):
            facultyitem = Faculty.objects.get(fId=f_id)
            if facultyitem.subjectCode != sCode:
                return HttpResponseForbidden("You are not authorized to update marks for this subject.")
            update = Student.objects.get(roll_no=roll, subjectCode=sCode)
            update.assignment1 = A1
            update.assignment2 = A2
            update.total = int(A1) + int(A2)
            update.save()

    facultyitem = Faculty.objects.get(fId=f_id)
    subCode = facultyitem.subjectCode
    allStudent = Student.objects.filter(subjectCode=subCode).order_by('roll_no')
    # for i in range(len(allStudent)):
    #     print(allStudent[i].classtest1)
    alldata = Student()
    st = Student.objects.all()
    return render(request, 'facultydashboard.html',
                  {'data': alldata, 'allStudent': allStudent, 'student': st, 'faculty': facultyitem})


def export_to_excel(request):
    # Ensure the faculty member is logged in
    if 'f_id' not in globals():
        return redirect('loginAsFaculty')  # Redirect to faculty login if not logged in

    # Get the logged-in faculty's subject code
    faculty = Faculty.objects.get(fId=f_id)
    faculty_subject_code = faculty.subjectCode

    # Fetch students with the matching subject code
    students = Student.objects.filter(subjectCode=faculty_subject_code).values(
        'roll_no', 'name', 'gender', 'parent', 'address', 'subjectCode', 'assignment1', 'assignment2'
    )

    # Create a DataFrame from the student data
    df = pd.DataFrame(list(students))
    df['total'] = df['assignment1'] + df['assignment2']  # Calculate the total

    # Create an Excel writer object
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=students.xlsx'

    # Convert DataFrame to Excel
    df.to_excel(response, index=False, engine='openpyxl')

    return response


def export_to_word(request):
    # Ensure the faculty member is logged in
    if 'f_id' not in globals():
        return redirect('loginAsFaculty')  # Redirect to faculty login if not logged in

    # Get the logged-in faculty's subject code
    faculty = Faculty.objects.get(fId=f_id)
    faculty_subject_code = faculty.subjectCode

    # Fetch students with the matching subject code
    students = Student.objects.filter(subjectCode=faculty_subject_code).values(
        'roll_no', 'name', 'gender', 'parent', 'address', 'subjectCode', 'assignment1', 'assignment2'
    )

    # Create a Document object
    document = Document()

    # Add a title to the document
    document.add_heading('Student Data Report', level=1)

    # Iterate over student data and add it to the document
    for student in students:
        roll_no = student['roll_no']
        name = student['name']
        gender = student['gender']
        parent = student['parent']
        address = student['address']
        subject_code = student['subjectCode']
        assignment1 = student['assignment1']
        assignment2 = student['assignment2']
        total = assignment1 + assignment2

        # Add student details to the document
        document.add_heading(f'Student Roll No: {roll_no}', level=2)
        document.add_paragraph(f'Name: {name}')
        document.add_paragraph(f'Gender: {gender}')
        document.add_paragraph(f'Parent: {parent}')
        document.add_paragraph(f'Address: {address}')
        document.add_paragraph(f'Subject Code: {subject_code}')
        document.add_paragraph(f'Assignment 1: {assignment1}')
        document.add_paragraph(f'Assignment 2: {assignment2}')
        document.add_paragraph(f'Total: {total}')
        document.add_paragraph('')  # Add a blank line for spacing

    # Create an HTTP response with the Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=students.docx'

    # Save the document to the response
    document.save(response)

    return response


def export_to_xml(request):
    # Ensure the faculty member is logged in
    if 'f_id' not in globals():
        return redirect('loginAsFaculty')  # Redirect to faculty login if not logged in

    # Get the logged-in faculty's subject code
    faculty = Faculty.objects.get(fId=f_id)
    faculty_subject_code = faculty.subjectCode

    # Fetch students with the matching subject code
    students = Student.objects.filter(subjectCode=faculty_subject_code).values(
        'roll_no', 'name', 'gender', 'parent', 'address', 'subjectCode', 'assignment1', 'assignment2'
    )

    # Create the root element
    root = ET.Element('Students')

    # Iterate over student data and add it to the XML
    for student in students:
        student_elem = ET.SubElement(root, 'Student')
        ET.SubElement(student_elem, 'RollNo').text = str(student['roll_no'])
        ET.SubElement(student_elem, 'Name').text = student['name']
        ET.SubElement(student_elem, 'Gender').text = student['gender']
        ET.SubElement(student_elem, 'Parent').text = student['parent']
        ET.SubElement(student_elem, 'Address').text = student['address']
        ET.SubElement(student_elem, 'SubjectCode').text = student['subjectCode']
        ET.SubElement(student_elem, 'Assignment1').text = str(student['assignment1'])
        ET.SubElement(student_elem, 'Assignment2').text = str(student['assignment2'])
        total = student['assignment1'] + student['assignment2']
        ET.SubElement(student_elem, 'Total').text = str(total)

    # Create a string from the XML tree
    xml_string = ET.tostring(root, encoding='utf-8')

    # Beautify the XML string
    pretty_xml = minidom.parseString(xml_string).toprettyxml(indent='  ')

    # Create an HTTP response with the XML content
    response = HttpResponse(pretty_xml, content_type='application/xml')
    response['Content-Disposition'] = 'attachment; filename=students.xml'

    return response


